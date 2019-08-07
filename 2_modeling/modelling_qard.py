import pandas as pd
import numpy as np
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from sklearn import model_selection
from sklearn import preprocessing
from imblearn.over_sampling import SMOTE
from imblearn.under_sampling import TomekLinks
from imblearn.combine import SMOTETomek
import matplotlib.pyplot as plt
from collections import Counter
from sklearn import tree
from sklearn import ensemble
from sklearn import metrics
from sklearn import naive_bayes
from sklearn import svm
from sklearn import neighbors
from sklearn import model_selection
import docx


def read_googlesheet(id_workbook, sheet_name):
    scope = ['https://spreadsheets.google.com/feeds']
    credentials = ServiceAccountCredentials.from_json_keyfile_name('qard-project-5db332216535.json', scope)
    gc = gspread.authorize(credentials)

    book = gc.open_by_key(id_workbook)
    worksheet = book.worksheet(sheet_name)
    table = worksheet.get_all_values()

    df = pd.DataFrame(table[1:], columns=table[0])
    return df


def prep_X_bool(X_train, X_test, X_bool_cols):
    if len(X_bool_cols)>0:
        X_bool_train = X_train[X_bool_cols].values
        X_bool_test = X_test[X_bool_cols].values
        feature_names = X_bool_cols
    else:
        X_bool_train = np.ones(len(X_train)).reshape(-1,1)
        X_bool_test = np.ones(len(X_test)).reshape(-1,1)
        feature_names = ['drop']
        
    return X_bool_train, X_bool_test, feature_names


def prep_X_cat(X_train, X_test, X_cat_cols):
    if len(X_cat_cols)>0:
        X_cat_train = X_train[X_cat_cols]
        X_cat_test = X_test[X_cat_cols]
        onehot = preprocessing.OneHotEncoder(sparse=False, handle_unknown='ignore')
        X_cat_train = onehot.fit_transform(X_cat_train)
        X_cat_test = onehot.transform(X_cat_test)
        feature_names = onehot.get_feature_names()
    else:
        X_cat_train = np.ones(len(X_train)).reshape(-1,1)
        X_cat_test = np.ones(len(X_test)).reshape(-1,1)
        feature_names = ['drop']
        
    return X_cat_train, X_cat_test, feature_names


def prep_y_num(X_train, X_test, X_num_cols, scaling):
    
    if len(X_num_cols)>0:
        
        X_num_train = X_train[X_num_cols]
        X_num_test = X_test[X_num_cols]
        feature_names = X_num_cols
        
        if scaling == 'std':
            # Standard scaling
            scaler = preprocessing.StandardScaler()
            X_num_train = scaler.fit_transform(X_num_train)
            X_num_test = scaler.transform(X_num_test)
        elif scaling == 'minmax':
            # Min-max scaling
            scaler = preprocessing.MinMaxScaler()
            X_num_train = scaler.fit_transform(X_num_train)
            X_num_test = scaler.transform(X_num_test)
        elif scaling == 'norm':
            # Normalise
            scaler = preprocessing.Normalizer()
            X_num_train = scaler.fit_transform(X_num_train)
            X_num_test = scaler.transform(X_num_test)
        elif scaling == 'none':
            X_num_train = X_num_train.values
            X_num_test = X_num_test.values
        else:
            raise ValueError('Normalisation method not recognized.')
    else:
        X_num_train = np.ones(len(X_train)).reshape(-1,1)
        X_num_test = np.ones(len(X_test)).reshape(-1,1)
        feature_names = ['drop']
        
    return X_num_train, X_num_test, feature_names


def resample(X, y, method):
    
    if method == 'smote':
        sm = SMOTE(random_state=2777, ratio=1.0)
        X, y = sm.fit_sample(X, y)
    elif method == 'tomek':
        tomek = TomekLinks(random_state=2777, sampling_strategy='majority')
        X, y = tomek.fit_sample(X, y)
    elif method == 'smote-tomek':
        smt = SMOTETomek(random_state=2777, ratio='auto')
        X, y = smt.fit_sample(X, y)
    elif method == 'none':
        pass
    else:
        raise ValueError('Resampling method not recognized.')
        
    return X, y


def prep_train_and_test(features_all, test_size=0.3, scaling='std', resampling='smote'):
    
    # Read Google sheet
    id_workbook = '1RVWLvccaKEAvufDuoQN6FUNBtco-MI5HSKoxQv-Ty3I'
    sheet_name = 'Features'
    features_info = read_googlesheet(id_workbook, sheet_name)
    
    # Filter features
    include = features_info[features_info['include'] == '1']
    included_features = include['name'].tolist()
    features = features_all[included_features]
    features = features.replace([np.inf, -np.inf], np.nan)
    features = features.dropna()

    #Create train test
    y = features['label']
    X = features.drop(columns=['label'])
    X_train, X_test, y_train, y_test = model_selection.train_test_split(X, y, test_size=test_size, random_state=46896)

    # Sorting of columns according to their categories
    X_num_cols = include.loc[include['type'] == 'numerical', 'name'].tolist()
    X_cat_cols = include.loc[include['type'] == 'categorical', 'name'].tolist()
    X_bool_cols = include.loc[include['type'] == 'boolean', 'name'].tolist()

    # Preparation of features 
    X_bool_train, X_bool_test, feat_bool = prep_X_bool(X_train, X_test, X_bool_cols)
    X_cat_train, X_cat_test, feat_cat = prep_X_cat(X_train, X_test, X_cat_cols)
    X_num_train, X_num_test, feat_num = prep_y_num(X_train, X_test, X_num_cols, scaling)

    # Concatenation of boolean, categorical and numeric features
    X_train_prep = np.concatenate((X_bool_train, X_cat_train, X_num_train), axis=1)
    X_test_prep = np.concatenate((X_bool_test, X_cat_test, X_num_test), axis=1)
    feature_names = np.concatenate((feat_bool, feat_cat, feat_num), axis=0)
    y_train_prep = y_train.values
    y_test_prep = y_test.values
    
    # Resampling
    X_train_prep, y_train_prep = resample(X_train_prep, y_train_prep, resampling)
    
    return X_train_prep, y_train_prep, X_test_prep, y_test_prep, feature_names

#----------------------------------------------------------------------------------------------------------

def predict_threshold(y_scores, threshold):
    y_pred = [1 if y > threshold else 0 for y in y_scores]
    return y_pred

def assess_model(X_train, y_train, X_test, y_test, feature_names, trained_model, threshold=0.5):
    # get y_scores
    y_train_scores = trained_model.predict_proba(X_train)[:, 1]
    y_test_scores = trained_model.predict_proba(X_test)[:, 1]
    
    # Prediction against train
    y_train_pred = predict_threshold(y_train_scores, threshold)
    cm_train = metrics.confusion_matrix(y_train, y_train_pred)
    precision_train = cm_train[0][0] / sum(cm_train[0])
    recall_train = cm_train[1][1] / sum(cm_train[1])

    # Performance against test
    y_test_pred = predict_threshold(y_test_scores, threshold)
    cm_test = metrics.confusion_matrix(y_test, y_test_pred)
    precision_test = cm_test[0][0] / sum(cm_test[0])
    recall_test = cm_test[1][1] / sum(cm_test[1])

    # ROC curve
    y_score_train = trained_model.predict_proba(X_train)
    y_score_test = trained_model.predict_proba(X_test)
    fpr_test, tpr_test, _ = metrics.roc_curve(y_test, y_score_test[:, 1])
    roc_auc_test = metrics.auc(fpr_test, tpr_test)
    fpr_train, tpr_train, _ = metrics.roc_curve(y_train, y_score_train[:, 1])
    roc_auc_train = metrics.auc(fpr_train, tpr_train)
    plt.figure(figsize=(10,10))
    lw=2
    plt.plot(fpr_test, tpr_test, color='red',
             lw=lw, label='ROC curve test (area = %0.2f)' % roc_auc_test)
    plt.plot(fpr_train, tpr_train, color='blue',
             lw=lw, label='ROC curve train (area = %0.2f)' % roc_auc_train)
    plt.plot([0, 1], [0, 1], color='navy', lw=lw, linestyle='--')
    plt.xlim([0.0, 1.0])
    plt.ylim([0.0, 1.05])
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('Receiver operating characteristic example')
    plt.legend(loc="lower right")
    plt.savefig('roc.png')
    
    # Print results to .docx
    try:
        document = docx.Document('models.docx')
    except:
        document = docx.Document()
    finally:
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(f'Recall on train: {recall_train}\nCorresponding confusion matrix:')
        r. add_break()
        r.add_text(f'{str(cm_train)}')
        r. add_break()
        r.add_text(f'Recall on test: {recall_test}\nCorresponding confusion matrix:')
        r. add_break()
        r.add_text(f'{str(cm_test)}')
        r. add_break()
        r.add_text(str(trained_model))
        r. add_break()
        r.add_text(f'Predictors: {feature_names}')
        r. add_break()
        r.add_text(f'Threshold: {threshold}')
        r. add_break()
        r.add_picture('roc.png', width=docx.shared.Inches(5.5))
        document.add_page_break()    
        document.save('models.docx')

    # Print results on screen
    print(f'Precision on train: {precision_train}')
    print(f'Recall on train: {recall_train}\nCorresponding confusion matrix:')
    print(f'{str(cm_train)}')
    print(f'Precision on test: {precision_test}')
    print(f'Recall on test: {recall_test}\nCorresponding confusion matrix:')
    print(f'{str(cm_test)}')
    print(f'Model: {str(trained_model)}')
    print(f'Predictors: {feature_names}')
    print(f'Threshold: {threshold}')
    plt.show()